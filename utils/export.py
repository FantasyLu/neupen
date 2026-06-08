"""
小说导出模块
支持将小说内容导出为 TXT、Markdown、Word 三种格式
"""

import os
import json
from datetime import datetime
from typing import Optional
from pathlib import Path

from core.models import get_db, Novel, Chapter, Volume
from core.config import DATA_DIR


# ======================================
# 导出工具类
# ======================================

class NovelExporter:
    """
    小说导出器
    从数据库读取小说内容，生成各种格式的文件
    """

    def __init__(self, novel_id: int):
        self.novel_id = novel_id
        self.db = get_db()
        self.export_dir = Path(DATA_DIR) / "exports"
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def _get_novel(self) -> Optional[Novel]:
        """获取小说基本信息"""
        return self.db.query(Novel).filter(Novel.id == self.novel_id).first()

    def _get_published_chapters(self) -> list[Chapter]:
        """获取所有已发布（有正文）的章节，按章节号排序"""
        return (
            self.db.query(Chapter)
            .filter(
                Chapter.novel_id == self.novel_id,
                Chapter.content.isnot(None),
                Chapter.content != ""
            )
            .order_by(Chapter.chapter_number)
            .all()
        )

    def _get_volumes(self) -> list[Volume]:
        """获取所有卷信息"""
        return (
            self.db.query(Volume)
            .filter(Volume.novel_id == self.novel_id)
            .order_by(Volume.volume_number)
            .all()
        )

    def _get_volume_for_chapter(self, chapter_number: int, volumes: list[Volume]) -> Optional[Volume]:
        """根据章节号找到所属卷"""
        for vol in volumes:
            if vol.start_chapter and vol.end_chapter:
                if vol.start_chapter <= chapter_number <= vol.end_chapter:
                    return vol
        return None

    def _build_filename(self, novel: Novel, ext: str) -> str:
        """生成导出文件名"""
        safe_title = "".join(c for c in novel.title if c.isalnum() or c in "._- 《》")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{safe_title}_{timestamp}.{ext}"

    # ======================================
    # TXT 导出
    # ======================================

    def export_txt(self, include_outline: bool = False) -> str:
        """
        导出为纯文本格式（TXT）

        Args:
            include_outline: 是否在正文前附上大纲摘要

        Returns:
            导出文件的绝对路径
        """
        novel = self._get_novel()
        if not novel:
            raise ValueError(f"小说项目 {self.novel_id} 不存在")

        chapters = self._get_published_chapters()
        if not chapters:
            raise ValueError("没有已完成的章节可供导出")

        volumes = self._get_volumes()
        lines = []

        # 封面信息
        lines.append(novel.title)
        lines.append("=" * 40)
        if novel.logline:
            lines.append(f"\n简介：{novel.logline}\n")
        if novel.genre:
            lines.append(f"题材：{novel.genre}")
        lines.append(f"总字数：{sum(ch.word_count or 0 for ch in chapters):,} 字")
        lines.append(f"章节数：{len(chapters)} 章")
        lines.append(f"导出时间：{datetime.now().strftime('%Y年%m月%d日')}")
        lines.append("\n" + "=" * 40 + "\n")

        # 大纲摘要（可选）
        if include_outline and novel.world_setting:
            try:
                world = json.loads(novel.world_setting)
                lines.append("【世界观设定】")
                for k, v in world.items():
                    lines.append(f"{k}：{v}")
                lines.append("\n" + "-" * 40 + "\n")
            except (json.JSONDecodeError, TypeError):
                pass

        # 正文内容（按卷分组）
        current_volume = None
        for chapter in chapters:
            vol = self._get_volume_for_chapter(chapter.chapter_number, volumes)
            if vol and (current_volume is None or vol.id != current_volume.id):
                current_volume = vol
                lines.append(f"\n{'=' * 40}")
                lines.append(f"第{vol.volume_number}卷  {vol.title}")
                if vol.summary:
                    lines.append(vol.summary)
                lines.append("=" * 40 + "\n")

            # 章节内容
            title = chapter.title or f"第{chapter.chapter_number}章"
            lines.append(f"\n第{chapter.chapter_number}章  {title}")
            lines.append("-" * 30)
            lines.append(chapter.content or "")
            lines.append("")  # 章节间空行

        content = "\n".join(lines)
        filename = self._build_filename(novel, "txt")
        filepath = self.export_dir / filename

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return str(filepath)

    # ======================================
    # Markdown 导出
    # ======================================

    def export_markdown(self, include_characters: bool = True,
                         include_foreshadowings: bool = False) -> str:
        """
        导出为 Markdown 格式

        Args:
            include_characters: 是否附上人物档案
            include_foreshadowings: 是否附上伏笔追踪

        Returns:
            导出文件的绝对路径
        """
        novel = self._get_novel()
        if not novel:
            raise ValueError(f"小说项目 {self.novel_id} 不存在")

        chapters = self._get_published_chapters()
        if not chapters:
            raise ValueError("没有已完成的章节可供导出")

        volumes = self._get_volumes()
        lines = []

        # YAML Front Matter（兼容 Obsidian 等 Markdown 编辑器）
        lines.append("---")
        lines.append(f"title: {novel.title}")
        lines.append(f"genre: {novel.genre or '未分类'}")
        lines.append(f"created: {novel.created_at.strftime('%Y-%m-%d') if novel.created_at else '未知'}")
        lines.append(f"exported: {datetime.now().strftime('%Y-%m-%d')}")
        lines.append(f"chapters: {len(chapters)}")
        lines.append(f"words: {sum(ch.word_count or 0 for ch in chapters)}")
        lines.append("---\n")

        # 标题和简介
        lines.append(f"# {novel.title}\n")
        if novel.logline:
            lines.append(f"> {novel.logline}\n")

        # 目录
        lines.append("## 目录\n")
        current_volume = None
        for chapter in chapters:
            vol = self._get_volume_for_chapter(chapter.chapter_number, volumes)
            if vol and (current_volume is None or vol.id != current_volume.id):
                current_volume = vol
                lines.append(f"\n**第{vol.volume_number}卷 · {vol.title}**\n")
            title = chapter.title or f"第{chapter.chapter_number}章"
            anchor = f"第{chapter.chapter_number}章{chapter.title or ''}"
            # Markdown 锚点（去掉特殊字符）
            anchor_id = anchor.replace(" ", "-").replace("《", "").replace("》", "")
            lines.append(f"- [第{chapter.chapter_number}章 {chapter.title or ''}](#{anchor_id})")

        lines.append("\n---\n")

        # 世界观设定（如果有）
        if novel.world_setting:
            try:
                world = json.loads(novel.world_setting)
                lines.append("## 世界观设定\n")
                for k, v in world.items():
                    lines.append(f"**{k}：** {v}\n")
                lines.append("\n---\n")
            except (json.JSONDecodeError, TypeError):
                pass

        # 人物档案（可选）
        if include_characters:
            from core.models import Character
            chars = (
                self.db.query(Character)
                .filter(Character.novel_id == self.novel_id)
                .order_by(Character.is_main.desc())
                .all()
            )
            if chars:
                lines.append("## 人物档案\n")
                for char in chars:
                    role_tag = "⭐ 主要人物" if char.is_main else "配角"
                    lines.append(f"### {char.name} `{char.role or role_tag}`\n")
                    if char.age or char.gender:
                        info = []
                        if char.gender: info.append(f"性别：{char.gender}")
                        if char.age: info.append(f"年龄：{char.age}")
                        lines.append(" | ".join(info) + "\n")
                    if char.personality:
                        lines.append(f"**性格：** {char.personality}\n")
                    if char.background:
                        lines.append(f"**背景：** {char.background}\n")
                    if char.motivations:
                        lines.append(f"**动机：** {char.motivations}\n")
                    if char.growth_arc:
                        lines.append(f"**成长弧光：** {char.growth_arc}\n")
                lines.append("\n---\n")

        # 正文内容（按卷分组）
        lines.append("## 正文\n")
        current_volume = None
        for chapter in chapters:
            vol = self._get_volume_for_chapter(chapter.chapter_number, volumes)
            if vol and (current_volume is None or vol.id != current_volume.id):
                current_volume = vol
                lines.append(f"## 第{vol.volume_number}卷 · {vol.title}\n")
                if vol.summary:
                    lines.append(f"> {vol.summary}\n")

            title_text = chapter.title or f"第{chapter.chapter_number}章"
            lines.append(f"### 第{chapter.chapter_number}章 {title_text}\n")

            # 章节元信息（可折叠）
            if chapter.word_count:
                lines.append(f"*字数：{chapter.word_count:,} 字*\n")

            lines.append(chapter.content or "")
            lines.append("\n---\n")

        # 伏笔追踪（可选）
        if include_foreshadowings:
            from core.models import Foreshadowing
            fs_list = (
                self.db.query(Foreshadowing)
                .filter(Foreshadowing.novel_id == self.novel_id)
                .order_by(Foreshadowing.set_chapter)
                .all()
            )
            if fs_list:
                lines.append("## 伏笔追踪\n")
                lines.append("| 伏笔名称 | 埋下章节 | 回收章节 | 状态 | 重要程度 |")
                lines.append("|---------|---------|---------|------|---------|")
                for fs in fs_list:
                    status_map = {"active": "⏳ 未回收", "collected": "✅ 已回收", "abandoned": "❌ 放弃"}
                    imp_map = {"high": "🔴 高", "medium": "🟡 中", "low": "🟢 低"}
                    lines.append(
                        f"| {fs.name} | 第{fs.set_chapter}章 | "
                        f"{'第' + str(fs.collect_chapter) + '章' if fs.collect_chapter else '-'} | "
                        f"{status_map.get(fs.status, fs.status)} | "
                        f"{imp_map.get(fs.importance, fs.importance)} |"
                    )

        content = "\n".join(lines)
        filename = self._build_filename(novel, "md")
        filepath = self.export_dir / filename

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return str(filepath)

    # ======================================
    # Word 导出
    # ======================================

    def export_word(self, include_characters: bool = True) -> str:
        """
        导出为 Word 格式（.docx）

        Args:
            include_characters: 是否附上人物档案

        Returns:
            导出文件的绝对路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            raise RuntimeError("请先安装 python-docx：pip install python-docx")

        novel = self._get_novel()
        if not novel:
            raise ValueError(f"小说项目 {self.novel_id} 不存在")

        chapters = self._get_published_chapters()
        if not chapters:
            raise ValueError("没有已完成的章节可供导出")

        volumes = self._get_volumes()
        doc = Document()

        # ---- 文档样式设置 ----
        # 设置页面边距
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        # ---- 封面页 ----
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(novel.title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = '黑体'

        doc.add_paragraph()  # 空行

        if novel.genre:
            genre_para = doc.add_paragraph()
            genre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            genre_para.add_run(f"【{novel.genre}】").font.size = Pt(14)

        if novel.logline:
            doc.add_paragraph()
            logline_para = doc.add_paragraph()
            logline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logline_run = logline_para.add_run(novel.logline)
            logline_run.font.size = Pt(12)
            logline_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()
        stats_para = doc.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_words = sum(ch.word_count or 0 for ch in chapters)
        stats_para.add_run(
            f"共 {len(chapters)} 章 · {total_words:,} 字\n"
            f"导出于 {datetime.now().strftime('%Y年%m月%d日')}"
        ).font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

        # ---- 人物档案（可选） ----
        if include_characters:
            from core.models import Character
            chars = (
                self.db.query(Character)
                .filter(Character.novel_id == self.novel_id)
                .order_by(Character.is_main.desc())
                .all()
            )
            if chars:
                heading = doc.add_heading("人物档案", level=1)
                heading.runs[0].font.name = '黑体'

                for char in chars:
                    char_heading = doc.add_heading(
                        f"{'★ ' if char.is_main else ''}{char.name}（{char.role or '配角'}）",
                        level=2
                    )
                    char_heading.runs[0].font.name = '黑体'

                    info_parts = []
                    if char.gender: info_parts.append(f"性别：{char.gender}")
                    if char.age: info_parts.append(f"年龄：{char.age}")
                    if info_parts:
                        doc.add_paragraph(" | ".join(info_parts))

                    if char.personality:
                        p = doc.add_paragraph()
                        p.add_run("性格特征：").bold = True
                        p.add_run(char.personality)
                    if char.background:
                        p = doc.add_paragraph()
                        p.add_run("背景故事：").bold = True
                        p.add_run(char.background)
                    if char.motivations:
                        p = doc.add_paragraph()
                        p.add_run("核心动机：").bold = True
                        p.add_run(char.motivations)
                    if char.growth_arc:
                        p = doc.add_paragraph()
                        p.add_run("成长弧光：").bold = True
                        p.add_run(char.growth_arc)

                doc.add_page_break()

        # ---- 正文内容 ----
        current_volume = None
        for chapter in chapters:
            vol = self._get_volume_for_chapter(chapter.chapter_number, volumes)

            # 卷标题（换页）
            if vol and (current_volume is None or vol.id != current_volume.id):
                current_volume = vol
                if chapter != chapters[0]:  # 不是第一卷时才换页
                    doc.add_page_break()
                vol_heading = doc.add_heading(
                    f"第{vol.volume_number}卷  {vol.title}", level=1
                )
                vol_heading.runs[0].font.name = '黑体'
                if vol.summary:
                    summary_para = doc.add_paragraph(vol.summary)
                    summary_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # 章节标题
            chapter_title = chapter.title or f"第{chapter.chapter_number}章"
            ch_heading = doc.add_heading(
                f"第{chapter.chapter_number}章  {chapter_title}", level=2
            )
            ch_heading.runs[0].font.name = '黑体'

            # 章节正文（保留段落结构）
            content = chapter.content or ""
            paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
            for para_text in paragraphs:
                p = doc.add_paragraph(para_text)
                # 对话段落缩进处理
                if para_text.startswith(("「", "『", "\u201c", '"', "【")):
                    p.paragraph_format.left_indent = Cm(0.5)
                else:
                    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符

        # ---- 保存文件 ----
        filename = self._build_filename(novel, "docx")
        filepath = self.export_dir / filename

        doc.save(str(filepath))
        return str(filepath)

    # ======================================
    # EPUB 导出
    # ======================================

    def export_epub(self, include_characters: bool = True,
                    include_foreshadowings: bool = False) -> str:
        """
        导出为 EPUB 电子书格式

        Args:
            include_characters: 是否附上人物档案
            include_foreshadowings: 是否附上伏笔追踪

        Returns:
            导出文件的绝对路径
        """
        try:
            from ebooklib import epub
        except ImportError:
            raise RuntimeError("请先安装 EbookLib：pip install EbookLib")

        novel = self._get_novel()
        if not novel:
            raise ValueError(f"小说项目 {self.novel_id} 不存在")

        chapters = self._get_published_chapters()
        if not chapters:
            raise ValueError("没有已完成的章节可供导出")

        volumes = self._get_volumes()

        book = epub.EpubBook()

        # ---- 元数据 ----
        book.set_identifier(f"neupen-novel-{self.novel_id}")
        book.set_title(novel.title)
        book.set_language("zh-CN")
        book.add_author(novel.author or "佚名")

        # ---- 全局 CSS ----
        css_content = """
body { font-family: "Noto Serif SC", "Source Han Serif CN", "宋体", serif;
       line-height: 1.8; }
h1 { text-align: center; margin: 2em 0 1em; }
h2 { margin: 1.5em 0 0.8em; }
h3 { margin: 1em 0 0.5em; }
p  { text-indent: 2em; margin: 0.4em 0; }
p.dialogue { text-indent: 0; padding-left: 2em; }
.meta { color: #999; font-size: 0.9em; text-align: center; }
.cover-title { font-size: 2em; font-weight: bold; text-align: center;
               margin-top: 30%; }
.cover-logline { text-align: center; color: #666; margin-top: 1em; }
.cover-stats { text-align: center; color: #999; margin-top: 2em; font-size: 0.9em; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #ccc; padding: 6px 10px; text-align: left; }
th { background: #f5f5f5; }
"""
        css = epub.EpubItem(uid="style", file_name="style/default.css",
                            media_type="text/css", content=css_content.encode("utf-8"))
        book.add_item(css)

        spine = ["nav"]
        toc = []
        epub_chapters = []

        # ---- 封面页 ----
        total_words = sum(ch.word_count or 0 for ch in chapters)
        cover_html = f"""<html><body>
<p class="cover-title">{novel.title}</p>
{'<p class="cover-logline">' + novel.logline + '</p>' if novel.logline else ''}
<p class="cover-stats">共 {len(chapters)} 章 · {total_words:,} 字<br/>
导出于 {datetime.now().strftime('%Y年%m月%d日')}</p>
</body></html>"""
        cover_ch = epub.EpubHtml(title="封面", file_name="cover.xhtml", lang="zh-CN")
        cover_ch.content = cover_html.encode("utf-8")
        cover_ch.add_item(css)
        book.add_item(cover_ch)
        spine.append(cover_ch)

        # ---- 人物档案（可选） ----
        if include_characters:
            from core.models import Character
            chars = (
                self.db.query(Character)
                .filter(Character.novel_id == self.novel_id)
                .order_by(Character.is_main.desc())
                .all()
            )
            if chars:
                char_parts = ["<h1>人物档案</h1>"]
                for char in chars:
                    role_tag = "⭐ 主要人物" if char.is_main else "配角"
                    char_parts.append(
                        f"<h2>{'★ ' if char.is_main else ''}{char.name}"
                        f"（{char.role or role_tag}）</h2>"
                    )
                    info = []
                    if char.gender: info.append(f"性别：{char.gender}")
                    if char.age: info.append(f"年龄：{char.age}")
                    if info:
                        char_parts.append(f"<p style='text-indent:0'>{' | '.join(info)}</p>")
                    if char.personality:
                        char_parts.append(f"<p><b>性格特征：</b>{char.personality}</p>")
                    if char.background:
                        char_parts.append(f"<p><b>背景故事：</b>{char.background}</p>")
                    if char.motivations:
                        char_parts.append(f"<p><b>核心动机：</b>{char.motivations}</p>")
                    if char.growth_arc:
                        char_parts.append(f"<p><b>成长弧光：</b>{char.growth_arc}</p>")

                char_ch = epub.EpubHtml(title="人物档案", file_name="characters.xhtml", lang="zh-CN")
                char_ch.content = ("<html><body>" + "\n".join(char_parts) + "</body></html>").encode("utf-8")
                char_ch.add_item(css)
                book.add_item(char_ch)
                spine.append(char_ch)
                toc.append(epub.Link("characters.xhtml", "人物档案", "characters"))

        # ---- 正文章节 ----
        current_volume = None
        vol_toc_children = []  # 当前卷下的章节列表
        vol_section = None

        def _flush_volume():
            """将上一个卷的 toc section 写入 toc"""
            nonlocal vol_section, vol_toc_children
            if vol_section and vol_toc_children:
                toc.append((vol_section, vol_toc_children))
            elif vol_toc_children:
                toc.extend(vol_toc_children)
            vol_toc_children = []
            vol_section = None

        for chapter in chapters:
            vol = self._get_volume_for_chapter(chapter.chapter_number, volumes)

            # 卷切换
            if vol and (current_volume is None or vol.id != current_volume.id):
                _flush_volume()
                current_volume = vol
                vol_section = epub.Section(f"第{vol.volume_number}卷 · {vol.title}")

            chapter_title = chapter.title or f"第{chapter.chapter_number}章"
            full_title = f"第{chapter.chapter_number}章  {chapter_title}"

            # 构建 HTML
            html_parts = [f"<h2>{full_title}</h2>"]
            if chapter.word_count:
                html_parts.append(f'<p class="meta">{chapter.word_count:,} 字</p>')

            content = chapter.content or ""
            for para in content.split("\n"):
                para = para.strip()
                if not para:
                    continue
                if para.startswith(("「", "『", "\u201c", '"', "【")):
                    html_parts.append(f'<p class="dialogue">{para}</p>')
                else:
                    html_parts.append(f"<p>{para}</p>")

            file_name = f"chapter_{chapter.chapter_number:04d}.xhtml"
            epub_ch = epub.EpubHtml(title=full_title, file_name=file_name, lang="zh-CN")
            epub_ch.content = ("<html><body>" + "\n".join(html_parts) + "</body></html>").encode("utf-8")
            epub_ch.add_item(css)
            book.add_item(epub_ch)
            epub_chapters.append(epub_ch)
            spine.append(epub_ch)
            vol_toc_children.append(epub.Link(file_name, full_title, f"ch{chapter.chapter_number}"))

        _flush_volume()

        # ---- 伏笔追踪（可选） ----
        if include_foreshadowings:
            from core.models import Foreshadowing
            fs_list = (
                self.db.query(Foreshadowing)
                .filter(Foreshadowing.novel_id == self.novel_id)
                .order_by(Foreshadowing.set_chapter)
                .all()
            )
            if fs_list:
                status_map = {"active": "未回收", "collected": "已回收", "abandoned": "放弃"}
                imp_map = {"high": "高", "medium": "中", "low": "低"}
                rows = []
                for fs in fs_list:
                    collect = f"第{fs.collect_chapter}章" if fs.collect_chapter else "-"
                    rows.append(
                        f"<tr><td>{fs.name}</td><td>第{fs.set_chapter}章</td>"
                        f"<td>{collect}</td>"
                        f"<td>{status_map.get(fs.status, fs.status)}</td>"
                        f"<td>{imp_map.get(fs.importance, fs.importance)}</td></tr>"
                    )
                fs_html = (
                    "<html><body><h1>伏笔追踪</h1>"
                    "<table><tr><th>伏笔名称</th><th>埋下章节</th><th>回收章节</th>"
                    "<th>状态</th><th>重要程度</th></tr>"
                    + "\n".join(rows) + "</table></body></html>"
                )
                fs_ch = epub.EpubHtml(title="伏笔追踪", file_name="foreshadowings.xhtml", lang="zh-CN")
                fs_ch.content = fs_html.encode("utf-8")
                fs_ch.add_item(css)
                book.add_item(fs_ch)
                spine.append(fs_ch)
                toc.append(epub.Link("foreshadowings.xhtml", "伏笔追踪", "foreshadowings"))

        # ---- 组装 ----
        book.toc = toc
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = spine

        # ---- 保存 ----
        filename = self._build_filename(novel, "epub")
        filepath = self.export_dir / filename
        epub.write_epub(str(filepath), book, {})

        return str(filepath)

    # ======================================
    # 便捷导出接口
    # ======================================

    def export(self, format: str = "txt", **kwargs) -> str:
        """
        统一导出接口

        Args:
            format: 导出格式，支持 'txt'、'markdown'/'md'、'word'/'docx'
            **kwargs: 传递给具体导出函数的参数

        Returns:
            导出文件的绝对路径
        """
        format = format.lower()
        if format == "txt":
            return self.export_txt(**kwargs)
        elif format in ("markdown", "md"):
            return self.export_markdown(**kwargs)
        elif format in ("word", "docx"):
            return self.export_word(**kwargs)
        elif format == "epub":
            return self.export_epub(**kwargs)
        else:
            raise ValueError(f"不支持的导出格式：{format}，请选择 txt/markdown/word/epub")

    def get_export_stats(self) -> dict:
        """
        获取当前可导出的统计信息
        """
        chapters = self._get_published_chapters()
        novel = self._get_novel()
        return {
            "novel_title": novel.title if novel else "未知",
            "published_chapters": len(chapters),
            "total_words": sum(ch.word_count or 0 for ch in chapters),
            "can_export": len(chapters) > 0
        }

    def close(self):
        """释放数据库连接"""
        self.db.close()


# ======================================
# 便捷函数
# ======================================

def export_novel(novel_id: int, format: str = "txt", **kwargs) -> str:
    """
    一键导出小说

    Args:
        novel_id: 小说项目 ID
        format: 导出格式 ('txt' / 'markdown' / 'word')
        **kwargs: 额外选项

    Returns:
        导出文件路径
    """
    exporter = NovelExporter(novel_id)
    try:
        return exporter.export(format, **kwargs)
    finally:
        exporter.close()
